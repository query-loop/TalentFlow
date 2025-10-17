"""ATS Formatter Agent

Generates ATS-friendly TXT and DOCX outputs from a structured profile and bullets.
Uses `python-docx` when available to create a .docx; otherwise returns a bytes
encoded placeholder for docx.
"""
from __future__ import annotations

from typing import Dict, List, Tuple
import io
import logging

try:
    from docx import Document
    _HAS_DOCX = True
except Exception:
    Document = None
    _HAS_DOCX = False

logger = logging.getLogger(__name__)


def format_txt(profile: Dict, bullets: List[str]) -> str:
    parts: List[str] = []
    parts.append(profile.get("name", ""))
    if profile.get("summary"):
        parts.append("\nSummary:\n" + profile.get("summary"))

    if bullets:
        parts.append("\nExperience:\n")
        parts.extend(bullets)

    if profile.get("skills"):
        sk = profile.get("skills")
        if isinstance(sk, list):
            parts.append("\nSkills:\n" + ", ".join(sk))

    if profile.get("education"):
        parts.append("\nEducation:\n")
        parts.extend(profile.get("education"))

    return "\n\n".join(parts)


def format_docx(profile: Dict, bullets: List[str]) -> bytes:
    # use python-docx if available
    txt = format_txt(profile, bullets)
    if _HAS_DOCX and Document is not None:
        doc = Document()
        doc.add_heading(profile.get("name", ""), level=1)
        if profile.get("summary"):
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(profile.get("summary"))
        if bullets:
            doc.add_heading("Experience", level=2)
            for b in bullets:
                doc.add_paragraph(b, style="List Bullet")
        if profile.get("skills"):
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(profile.get("skills")))
        if profile.get("education"):
            doc.add_heading("Education", level=2)
            for e in profile.get("education"):
                doc.add_paragraph(e)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # fallback: return TXT bytes as placeholder
    return txt.encode("utf-8")


def format_outputs(profile: Dict, bullets: List[str]) -> Tuple[str, bytes]:
    txt = format_txt(profile, bullets)
    docx = format_docx(profile, bullets)
    return txt, docx
